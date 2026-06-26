from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")


def fill_resume_template(template_path: Path, output_path: Path, content_map: dict[str, str]) -> dict[str, Any]:
    """按 run 级占位符替换填充 DOCX 简历模板，保留原 run 样式。"""
    document = Document(str(template_path))
    replaced: list[str] = []
    for paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            original = run.text
            filled = replace_placeholders(original, content_map, replaced)
            if filled != original:
                run.text = filled
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    unique_replaced = list(dict.fromkeys(replaced))
    unused_keys = [key for key in content_map if key not in unique_replaced]
    return {
        "status": "SUCCEEDED",
        "outputPath": str(output_path),
        "placeholders": unique_replaced,
        "unusedKeys": unused_keys,
    }


def replace_placeholders(text: str, content_map: dict[str, str], replaced: list[str]) -> str:
    """替换单个 run 文本中的所有占位符。"""

    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        if key not in content_map:
            return match.group(0)
        replaced.append(key)
        return str(content_map[key])

    return PLACEHOLDER_PATTERN.sub(replace, text)


def iter_paragraphs(document: DocxDocument):
    """遍历正文和表格单元格中的段落。"""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table: Table):
    """遍历表格内所有单元格段落。"""
    for row in table.rows:
        for cell in row.cells:
            yield from iter_cell_paragraphs(cell)


def iter_cell_paragraphs(cell: _Cell):
    """遍历单元格段落和嵌套表格。"""
    for paragraph in cell.paragraphs:
        yield paragraph
    for nested_table in cell.tables:
        yield from iter_table_paragraphs(nested_table)
