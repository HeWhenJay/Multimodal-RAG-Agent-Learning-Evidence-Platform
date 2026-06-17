from __future__ import annotations

import os
import shlex
import subprocess
import tempfile
from pathlib import Path

from rag.models import ParsedDocument


class MineruDocumentLoader:
    """MinerU adapter with a deterministic local fallback for development."""

    def __init__(self, command: str | None = None) -> None:
        self.command = command or os.getenv("MINERU_COMMAND")

    def load_bytes(self, content: bytes, filename: str, content_type: str | None = None) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if self.command:
            parsed = self._load_with_mineru(content, filename)
            if parsed.text.strip():
                return parsed

        text = self._fallback_parse(content, suffix)
        return ParsedDocument(
            text=text,
            parser="fallback-local-parser",
            metadata={"filename": filename, "contentType": content_type, "suffix": suffix},
        )

    def _load_with_mineru(self, content: bytes, filename: str) -> ParsedDocument:
        with tempfile.TemporaryDirectory(prefix="mineru-rag-") as tmp:
            tmp_dir = Path(tmp)
            input_path = tmp_dir / filename
            output_dir = tmp_dir / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            input_path.write_bytes(content)

            command = self.command or ""
            if "{input}" in command or "{output}" in command:
                command = command.format(input=str(input_path), output=str(output_dir))
                args = shlex.split(command)
            else:
                args = shlex.split(command) + [str(input_path), str(output_dir)]

            try:
                subprocess.run(args, check=True, timeout=120, capture_output=True, text=True)
            except Exception:
                return ParsedDocument(text="", parser="mineru-failed", metadata={"filename": filename})

            for path in list(output_dir.rglob("*.md")) + list(output_dir.rglob("*.txt")):
                text = path.read_text(encoding="utf-8", errors="ignore")
                if text.strip():
                    return ParsedDocument(text=text, parser="mineru", metadata={"filename": filename})

        return ParsedDocument(text="", parser="mineru-empty", metadata={"filename": filename})

    def _fallback_parse(self, content: bytes, suffix: str) -> str:
        if suffix == ".pdf":
            text = self._parse_pdf(content)
            if text.strip():
                return text
        if suffix == ".docx":
            text = self._parse_docx(content)
            if text.strip():
                return text
        return content.decode("utf-8", errors="ignore")

    def _parse_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            reader = PdfReader(tmp_path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""
        finally:
            if "tmp_path" in locals():
                Path(tmp_path).unlink(missing_ok=True)

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            document = Document(tmp_path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return ""
        finally:
            if "tmp_path" in locals():
                Path(tmp_path).unlink(missing_ok=True)

