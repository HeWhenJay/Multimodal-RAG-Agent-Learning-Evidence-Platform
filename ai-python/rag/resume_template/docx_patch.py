from __future__ import annotations

import base64
import copy
import hashlib
import io
import json
import re
import uuid
import zipfile
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from lxml import etree

from app.schemas.resume_template import (
    EvidencePolicy,
    LayoutChangeContract,
    ResumeContentPatch,
    ResumeTemplateBinding,
    ResumeTemplateExportResponse,
    ResumeTemplateLocationRef,
    ResumeTemplateParseResponse,
    ResumePatchValidationResponse,
    SectionKey,
)


FORBIDDEN_PATCH_KEYS = {
    "style",
    "font",
    "layout",
    "xml",
    "path",
    "location",
    "locationref",
    "locationrefs",
    "run",
    "runs",
    "paragraph",
    "table",
    "cell",
    "header",
    "footer",
}
INJECTION_PATTERNS = [
    re.compile(r"<\/?(?:w:|xml|html|body|table|tr|td|p|div|span|script|style)[^>]*>", re.IGNORECASE),
    re.compile(r"\|[^\n]+\|[ \t]*\n[ \t]*\|[ \t:.-]+\|"),
    re.compile(r"<\?xml", re.IGNORECASE),
    re.compile(r"<w:", re.IGNORECASE),
]
UNSUPPORTED_TAGS = {
    "hyperlink": "超链接",
    "fldChar": "域代码",
    "instrText": "域代码",
    "commentRangeStart": "批注",
    "commentRangeEnd": "批注",
    "ins": "修订痕迹",
    "del": "修订痕迹",
    "footnoteReference": "脚注",
    "endnoteReference": "尾注",
    "drawing": "绘图或图片",
    "pict": "旧版图片或形状",
    "txbxContent": "文本框",
    "smartTag": "智能标记",
    "sdt": "内容控件",
}
EXPORT_BLOCKING_TAGS = {
    local_name: label
    for local_name, label in UNSUPPORTED_TAGS.items()
    if local_name not in {"drawing", "pict"}
}
SECTION_KEYWORDS: list[tuple[SectionKey, tuple[str, ...]]] = [
    ("personal_info", ("姓名", "电话", "邮箱", "求职意向", "个人信息")),
    ("summary", ("个人总结", "自我评价", "优势", "简介")),
    ("education", ("教育", "学校", "学院", "专业", "学位")),
    ("work_experience", ("工作", "实习", "任职", "公司")),
    ("project_experience", ("项目", "作品", "实践", "系统")),
    ("skills", ("技能", "技术", "语言", "工具", "栈")),
    ("awards", ("奖项", "荣誉", "竞赛")),
    ("certifications", ("证书", "认证")),
    ("research", ("论文", "科研", "研究")),
]


@dataclass
class FieldTarget:
    """保存一次可应用替换的内存定位。"""

    field_id: str
    paragraph: Paragraph
    location: ResumeTemplateLocationRef


def parse_resume_template_docx(content: bytes, filename: str, template_id: str | None = None, version: int = 1) -> ResumeTemplateParseResponse:
    """解析 DOCX，生成字段绑定、文本 hash 和版式 fingerprint。"""
    if not content:
        raise ValueError("简历模板文件不能为空")
    if not _looks_like_docx(content):
        raise ValueError("当前只支持 DOCX 简历模板解析")
    template_id = template_id or uuid.uuid4().hex
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"DOCX 文件无法被 python-docx 读取：{exc}") from exc
    unsupported = detect_unsupported_regions(content)
    fields: list[ResumeTemplateBinding] = []
    for target in _iter_field_targets(document, template_id, version):
        source_text = _paragraph_text(target.paragraph)
        if not source_text.strip():
            continue
        local_unsupported = _paragraph_unsupported_reasons(target.paragraph)
        binding = ResumeTemplateBinding(
            templateId=template_id,
            version=version,
            fieldId=target.field_id,
            sectionKey=infer_section_key(source_text),
            displayName=build_display_name(source_text, len(fields) + 1),
            sourceText=source_text,
            sourceTextHash=hash_text(source_text),
            locationRefs=[target.location],
            styleFingerprint=paragraph_style_fingerprint(target.paragraph),
            maxChars=max(80, min(800, len(source_text) + 120)),
            maxLines=max(1, min(8, source_text.count("\n") + 2)),
            requiredEvidencePolicy=required_evidence_policy(infer_section_key(source_text)),
            unsupportedRegions=local_unsupported,
        )
        fields.append(binding)
    return ResumeTemplateParseResponse(
        templateId=template_id,
        version=version,
        filename=filename,
        fields=fields,
        unsupportedRegions=unsupported,
        layoutFingerprint=safe_layout_fingerprint(content),
    )


def validate_resume_patches(
    template_id: str,
    version: int,
    fields: list[ResumeTemplateBinding],
    patches: list[ResumeContentPatch],
    allowed_evidence_ids: list[str] | None = None,
    layout_contract: LayoutChangeContract | None = None,
) -> ResumePatchValidationResponse:
    """校验字段补丁是否只包含可接受的内容变更。"""
    contract = layout_contract or LayoutChangeContract()
    allowed_evidence = set(allowed_evidence_ids or [])
    fields_by_id = {field.fieldId: field for field in fields}
    errors: list[str] = []
    valid_patches: list[ResumeContentPatch] = []
    seen: set[str] = set()
    for patch in patches:
        path = f"patch[{patch.fieldId}]"
        raw = patch.model_dump()
        forbidden_keys = sorted(key for key in raw.keys() if key.lower() in FORBIDDEN_PATCH_KEYS)
        if forbidden_keys:
            errors.append(f"{path} 包含禁止的排版字段: {', '.join(forbidden_keys)}")
        field = fields_by_id.get(patch.fieldId)
        if field is None:
            errors.append(f"{path} 不存在对应字段")
            continue
        if field.templateId != template_id or field.version != version:
            errors.append(f"{path} 模板版本不匹配")
        if patch.fieldId in seen:
            errors.append(f"{path} 重复提交")
        seen.add(patch.fieldId)
        if patch.sourceTextHash != field.sourceTextHash:
            errors.append(f"{path} 原文字段 hash 不匹配")
        if field.unsupportedRegions:
            errors.append(f"{path} 位于暂不支持区域，拒绝自动修改")
        text_errors = validate_patch_text(patch.newText, field.maxChars, field.maxLines)
        errors.extend(f"{path} {item}" for item in text_errors)
        if field.requiredEvidencePolicy == "REQUIRED" and not patch.evidenceIds:
            errors.append(f"{path} 缺少必需 evidence 引用")
        invalid_evidence = [evidence_id for evidence_id in patch.evidenceIds if allowed_evidence and evidence_id not in allowed_evidence]
        if invalid_evidence:
            errors.append(f"{path} evidenceIds 不属于候选集合: {', '.join(invalid_evidence)}")
        if any(flag == "NONE" for flag in patch.riskFlags) and len(patch.riskFlags) > 1:
            errors.append(f"{path} riskFlags 中 NONE 不能和其他风险同时出现")
        if not any(error.startswith(path) for error in errors):
            valid_patches.append(patch)
    errors.extend(validate_layout_contract_for_patches(contract, valid_patches, fields_by_id))
    return ResumePatchValidationResponse(
        templateId=template_id,
        version=version,
        patches=valid_patches,
        validationErrors=errors,
        layoutValidation={
            "status": "PASSED" if not errors else "FAILED",
            "mode": contract.mode,
            "allowedChanges": [change.model_dump() for change in contract.allowedChanges],
            "message": "版式变更契约校验通过" if not errors else "版式变更契约校验失败",
        },
    )


def apply_resume_patches_to_docx(
    content: bytes,
    filename: str,
    template_id: str,
    version: int,
    fields: list[ResumeTemplateBinding],
    patches: list[ResumeContentPatch],
    allowed_evidence_ids: list[str] | None = None,
    layout_contract: LayoutChangeContract | None = None,
) -> ResumeTemplateExportResponse:
    """把已确认补丁确定性应用到 DOCX，并校验版式 fingerprint 不变。"""
    contract = layout_contract or LayoutChangeContract()
    blocking_regions = detect_export_blocking_regions(content)
    if blocking_regions:
        raise ValueError("简历包含暂不支持的复杂结构，拒绝自动导出：" + "；".join(blocking_regions))
    validation = validate_resume_patches(template_id, version, fields, patches, allowed_evidence_ids, contract)
    if validation.validationErrors:
        raise ValueError("补丁校验失败: " + "；".join(validation.validationErrors))
    original_fingerprint = layout_fingerprint(content)
    document = Document(io.BytesIO(content))
    patches_by_id = {patch.fieldId: patch for patch in validation.patches if patch.status in {"CONFIRMED", "VALIDATED"}}
    applied = 0
    for field in fields:
        patch = patches_by_id.get(field.fieldId)
        if patch is None:
            continue
        if len(field.locationRefs) != 1:
            raise ValueError(f"字段 {field.fieldId} 包含多个定位，当前版本拒绝自动修改")
        location = field.locationRefs[0]
        paragraph = resolve_paragraph(document, location)
        current_text = _paragraph_text(paragraph)
        if hash_text(current_text) != field.sourceTextHash:
            raise ValueError(f"字段 {field.fieldId} 当前文本 hash 已变化，拒绝导出")
        replace_paragraph_text_in_existing_runs(paragraph, patch.newText)
        applied += 1
    output = io.BytesIO()
    document.save(output)
    output_bytes = output.getvalue()
    new_fingerprint = layout_fingerprint(output_bytes)
    layout_validation = validate_layout_change(original_fingerprint, new_fingerprint, contract)
    if layout_validation["status"] != "PASSED":
        raise ValueError("RESUME_LAYOUT_CHANGED")
    return ResumeTemplateExportResponse(
        templateId=template_id,
        version=version + 1,
        filename=filename,
        fileBase64=base64.b64encode(output_bytes).decode("ascii"),
        layoutValidation=layout_validation,
        appliedPatchCount=applied,
    )


def validate_layout_contract_for_patches(
    contract: LayoutChangeContract,
    patches: list[ResumeContentPatch],
    fields_by_id: dict[str, ResumeTemplateBinding],
) -> list[str]:
    """校验补丁意图是否符合用户授权的版式变更契约。"""
    errors: list[str] = []
    if contract.mode == "PRESERVE_LAYOUT":
        layout_risk_patches = [patch.fieldId for patch in patches if "LAYOUT_RISK" in patch.riskFlags]
        if layout_risk_patches:
            errors.append("PRESERVE_LAYOUT 模式不允许带 LAYOUT_RISK 的补丁: " + "、".join(layout_risk_patches))
        return errors
    if contract.mode == "CONTROLLED_EDIT":
        allowed_field_ids = {change.fieldId for change in contract.allowedChanges if change.fieldId}
        for change in contract.allowedChanges:
            if change.fieldId and change.fieldId not in fields_by_id:
                errors.append(f"LayoutChangeContract 指向未知字段: {change.fieldId}")
            if change.type == "STYLE_RANGE" and not change.stylePatch:
                errors.append(f"STYLE_RANGE 必须声明 stylePatch: {change.fieldId or change.sectionKey or '未指定字段'}")
        for patch in patches:
            if "LAYOUT_RISK" in patch.riskFlags and patch.fieldId not in allowed_field_ids:
                errors.append(f"字段 {patch.fieldId} 标记 LAYOUT_RISK，但未在 LayoutChangeContract 中授权")
        return errors
    return errors


def validate_layout_change(
    baseline: dict[str, Any],
    current: dict[str, Any],
    contract: LayoutChangeContract,
) -> dict[str, Any]:
    """根据版式变更契约审计 DOCX 结构 fingerprint 差异。"""
    diff = layout_fingerprint_diff(baseline, current)
    if contract.mode == "PRESERVE_LAYOUT":
        passed = not diff["changedKeys"]
        return {
            "status": "PASSED" if passed else "FAILED",
            "mode": contract.mode,
            "baseline": baseline,
            "current": current,
            "diff": diff,
            "message": "XML 结构 fingerprint 未变化" if passed else "RESUME_LAYOUT_CHANGED: 存在未授权版式变化",
        }
    if contract.mode == "CONTROLLED_EDIT":
        blocking_keys = [
            key for key in diff["changedKeys"]
            if key not in {"paragraphCount", "runCount", "structureHash"}
        ]
        paragraph_delta = abs(int(current.get("paragraphCount", 0)) - int(baseline.get("paragraphCount", 0)))
        run_delta = abs(int(current.get("runCount", 0)) - int(baseline.get("runCount", 0)))
        errors: list[str] = []
        if blocking_keys:
            errors.append("存在未授权结构变化: " + "、".join(blocking_keys))
        if paragraph_delta > contract.maxParagraphDelta:
            errors.append(f"段落变化 {paragraph_delta} 超过授权上限 {contract.maxParagraphDelta}")
        if run_delta > contract.maxRunDelta:
            errors.append(f"run 变化 {run_delta} 超过授权上限 {contract.maxRunDelta}")
        if diff["changedKeys"] and not contract.allowedChanges:
            errors.append("检测到结构变化，但 LayoutChangeContract 未声明 allowedChanges")
        return {
            "status": "PASSED" if not errors else "FAILED",
            "mode": contract.mode,
            "baseline": baseline,
            "current": current,
            "diff": diff,
            "allowedChanges": [change.model_dump() for change in contract.allowedChanges],
            "message": "授权版式变化已通过结构审计" if not errors else "RESUME_LAYOUT_CHANGED: " + "；".join(errors),
        }
    return {
        "status": "REVIEW_REQUIRED",
        "mode": contract.mode,
        "baseline": baseline,
        "current": current,
        "diff": diff,
        "message": "RELAYOUT 模式需要人工预览确认后保存",
    }


def layout_fingerprint_diff(baseline: dict[str, Any], current: dict[str, Any]) -> dict[str, Any]:
    """生成两个版式指纹的可解释差异。"""
    keys = sorted(set(baseline.keys()) | set(current.keys()))
    changed = [key for key in keys if baseline.get(key) != current.get(key)]
    return {
        "changedKeys": changed,
        "paragraphDelta": int(current.get("paragraphCount", 0)) - int(baseline.get("paragraphCount", 0)),
        "tableDelta": int(current.get("tableCount", 0)) - int(baseline.get("tableCount", 0)),
        "runDelta": int(current.get("runCount", 0)) - int(baseline.get("runCount", 0)),
        "mediaChanged": baseline.get("mediaNames") != current.get("mediaNames"),
        "relationshipsChanged": baseline.get("relationshipHashes") != current.get("relationshipHashes"),
        "structureHashChanged": baseline.get("structureHash") != current.get("structureHash"),
    }


def validate_patch_text(new_text: str, max_chars: int, max_lines: int) -> list[str]:
    """校验补丁文本长度和注入风险。"""
    errors: list[str] = []
    if len(new_text) > max_chars:
        errors.append(f"超过字段长度限制 {max_chars}")
    if new_text.count("\n") + 1 > max_lines:
        errors.append(f"超过字段行数限制 {max_lines}")
    if any(pattern.search(new_text) for pattern in INJECTION_PATTERNS):
        errors.append("包含 Markdown 表格、HTML 或 DOCX XML 注入风险")
    return errors


def replace_paragraph_text_in_existing_runs(paragraph: Paragraph, new_text: str) -> None:
    """只改写既有 w:r/w:t 文本节点，不新增或删除 run。"""
    text_nodes = paragraph._element.xpath(".//w:r/w:t")
    if not text_nodes:
        raise ValueError("字段没有可替换的 w:t 文本节点")
    if len(text_nodes) > 1:
        raise ValueError("字段跨多个 run，当前版本拒绝自动修改")
    text_nodes[0].text = new_text


def resolve_paragraph(document: DocxDocument, location: ResumeTemplateLocationRef) -> Paragraph:
    """根据解析器生成的 locationRef 重新定位段落。"""
    if location.partName != "document":
        raise ValueError("当前版本只支持正文区域字段")
    if location.containerType == "paragraph":
        paragraphs = list(document.paragraphs)
        if location.paragraphIndex >= len(paragraphs):
            raise ValueError("字段段落定位已失效")
        return paragraphs[location.paragraphIndex]
    if location.containerType == "table_cell":
        tables = list(document.tables)
        if location.tableIndex is None or location.tableIndex >= len(tables):
            raise ValueError("字段表格定位已失效")
        table = tables[location.tableIndex]
        if location.rowIndex is None or location.rowIndex >= len(table.rows):
            raise ValueError("字段表格行定位已失效")
        row = table.rows[location.rowIndex]
        if location.cellIndex is None or location.cellIndex >= len(row.cells):
            raise ValueError("字段表格单元格定位已失效")
        cell = row.cells[location.cellIndex]
        if location.paragraphIndex >= len(cell.paragraphs):
            raise ValueError("字段单元格段落定位已失效")
        return cell.paragraphs[location.paragraphIndex]
    raise ValueError("不支持的字段定位类型")


def layout_fingerprint(content: bytes) -> dict[str, Any]:
    """生成不包含正文文本的 DOCX 结构 fingerprint。"""
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        names = sorted(archive.namelist())
        rel_hashes = {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in names
            if name.startswith("word/_rels/") or name == "word/_rels/document.xml.rels"
        }
        media_names = [name for name in names if name.startswith("word/media/")]
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    nsmap = {key: value for key, value in root.nsmap.items() if key}
    structure_root = copy.deepcopy(root)
    for node in structure_root.xpath(".//w:t", namespaces=nsmap):
        node.text = ""
    paragraphs = structure_root.xpath(".//w:body//w:p", namespaces=nsmap)
    tables = structure_root.xpath(".//w:body//w:tbl", namespaces=nsmap)
    runs = structure_root.xpath(".//w:body//w:r", namespaces=nsmap)
    structure_hash = hashlib.sha256(etree.tostring(structure_root, with_tail=False)).hexdigest()
    return {
        "paragraphCount": len(paragraphs),
        "tableCount": len(tables),
        "runCount": len(runs),
        "mediaNames": media_names,
        "relationshipHashes": rel_hashes,
        "structureHash": structure_hash,
    }


def safe_layout_fingerprint(content: bytes) -> dict[str, Any]:
    """生成版式指纹，异常时返回可诊断降级摘要而不中断上传解析。"""
    try:
        return layout_fingerprint(content)
    except Exception as exc:
        return {
            "paragraphCount": 0,
            "tableCount": 0,
            "runCount": 0,
            "mediaNames": [],
            "relationshipHashes": {},
            "structureHash": "unavailable",
            "warning": f"版式指纹生成失败：{exc}",
        }


def detect_unsupported_regions(content: bytes) -> list[str]:
    """扫描 DOCX XML 中首版不支持自动修改的复杂结构。"""
    return _detect_regions_by_tags(content, UNSUPPORTED_TAGS)


def detect_export_blocking_regions(content: bytes) -> list[str]:
    """扫描导出阶段必须全局拒绝的复杂结构，静态图片由字段级校验保护。"""
    return _detect_regions_by_tags(content, EXPORT_BLOCKING_TAGS)


def _detect_regions_by_tags(content: bytes, tag_labels: dict[str, str]) -> list[str]:
    """按指定标签集合扫描 DOCX XML 结构。"""
    unsupported: list[str] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            xml = archive.read(name)
            try:
                root = etree.fromstring(xml)
            except etree.XMLSyntaxError:
                continue
            for local_name, label in tag_labels.items():
                if root.xpath(f".//*[local-name()='{local_name}']"):
                    unsupported.append(f"{name}: {label}")
    return sorted(set(unsupported))


def _iter_field_targets(document: DocxDocument, template_id: str, version: int) -> list[FieldTarget]:
    """遍历普通正文段落和表格单元格，生成字段定位。"""
    targets: list[FieldTarget] = []
    for index, paragraph in enumerate(document.paragraphs):
        try:
            if _is_heading(paragraph) or _paragraph_unsupported_reasons(paragraph):
                continue
            if _paragraph_text(paragraph).strip():
                location = ResumeTemplateLocationRef(
                    partName="document",
                    containerType="paragraph",
                    paragraphIndex=index,
                    runStart=0,
                    runEnd=max(0, len(paragraph.runs) - 1),
                    textStart=0,
                    textEnd=len(_paragraph_text(paragraph)),
                )
                targets.append(FieldTarget(build_field_id(template_id, version, "p", len(targets) + 1), paragraph, location))
        except Exception:
            continue
    for table_index, table in enumerate(document.tables):
        try:
            targets.extend(_iter_table_field_targets(table, template_id, version, table_index, len(targets)))
        except Exception:
            continue
    return targets


def _iter_table_field_targets(table: Table, template_id: str, version: int, table_index: int, offset: int) -> list[FieldTarget]:
    """遍历普通表格单元格段落。"""
    targets: list[FieldTarget] = []
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            targets.extend(_iter_cell_field_targets(cell, template_id, version, table_index, row_index, cell_index, offset + len(targets)))
    return targets


def _iter_cell_field_targets(
    cell: _Cell,
    template_id: str,
    version: int,
    table_index: int,
    row_index: int,
    cell_index: int,
    offset: int,
) -> list[FieldTarget]:
    """遍历单元格内可修改段落。"""
    targets: list[FieldTarget] = []
    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        try:
            if _paragraph_unsupported_reasons(paragraph):
                continue
            source_text = _paragraph_text(paragraph)
            if not source_text.strip():
                continue
            location = ResumeTemplateLocationRef(
                partName="document",
                containerType="table_cell",
                paragraphIndex=paragraph_index,
                tableIndex=table_index,
                rowIndex=row_index,
                cellIndex=cell_index,
                runStart=0,
                runEnd=max(0, len(paragraph.runs) - 1),
                textStart=0,
                textEnd=len(source_text),
            )
            targets.append(FieldTarget(build_field_id(template_id, version, "tc", offset + len(targets) + 1), paragraph, location))
        except Exception:
            continue
    return targets


def build_field_id(template_id: str, version: int, prefix: str, index: int) -> str:
    """生成稳定字段 ID，避免暴露 XML 路径。"""
    return f"{prefix}-{hashlib.sha1(f'{template_id}:{version}:{prefix}:{index}'.encode('utf-8')).hexdigest()[:10]}"


def infer_section_key(text: str) -> SectionKey:
    """根据字段内容粗略推断简历区块类型。"""
    normalized = text.lower()
    for section, keywords in SECTION_KEYWORDS:
        if any(keyword.lower() in normalized for keyword in keywords):
            return section
    return "other"


def required_evidence_policy(section_key: SectionKey) -> EvidencePolicy:
    """项目、工作、技能等经历类字段默认要求 evidence 支撑。"""
    return "REQUIRED" if section_key in {"work_experience", "project_experience", "skills", "research"} else "OPTIONAL"


def build_display_name(source_text: str, index: int) -> str:
    """生成前端可读的字段名。"""
    cleaned = re.sub(r"\s+", " ", source_text).strip()
    return cleaned[:24] or f"简历字段 {index}"


def paragraph_style_fingerprint(paragraph: Paragraph) -> dict[str, Any]:
    """记录段落和 run 样式摘要，用于前端展示和安全校验。"""
    return {
        "styleName": safe_style_name(paragraph),
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "runCount": len(paragraph.runs),
        "runs": [
            {
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline),
                "styleName": safe_run_style_name(run),
            }
            for run in paragraph.runs
        ],
    }


def _paragraph_text(paragraph: Paragraph) -> str:
    """读取段落 run 文本，不使用 paragraph.text 写回。"""
    return "".join(run.text for run in paragraph.runs)


def _paragraph_unsupported_reasons(paragraph: Paragraph) -> list[str]:
    """判断单个字段是否含有首版不支持的复杂结构。"""
    reasons: list[str] = []
    for local_name, label in UNSUPPORTED_TAGS.items():
        if paragraph._element.xpath(f".//*[local-name()='{local_name}']"):
            reasons.append(label)
    if len(paragraph.runs) != 1 and _paragraph_text(paragraph).strip():
        reasons.append("跨 run 字段")
    return sorted(set(reasons))


def _is_heading(paragraph: Paragraph) -> bool:
    """标题段落只用于上下文，不作为可直接改写字段。"""
    style_name = safe_style_name(paragraph).lower()
    return style_name.startswith("heading") or style_name.startswith("标题")


def safe_style_name(paragraph: Paragraph) -> str:
    """读取段落样式名，样式缺失或损坏时返回空字符串。"""
    try:
        style = paragraph.style
        return style.name if style is not None and style.name else ""
    except Exception:
        return ""


def safe_run_style_name(run: Any) -> str:
    """读取 run 样式名，样式缺失或损坏时返回空字符串。"""
    try:
        style = run.style
        return style.name if style is not None and style.name else ""
    except Exception:
        return ""


def hash_text(text: str) -> str:
    """计算字段原文 hash。"""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _looks_like_docx(content: bytes) -> bool:
    """检查上传内容是否包含 DOCX 主文档。"""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            return "word/document.xml" in archive.namelist()
    except zipfile.BadZipFile:
        return False


def canonical_json(value: Any) -> str:
    """生成稳定 JSON 字符串，便于 Java/Python 之间存储。"""
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
