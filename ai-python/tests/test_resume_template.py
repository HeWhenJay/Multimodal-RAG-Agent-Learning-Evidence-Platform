from pathlib import Path

from docx import Document

from agent.resume_template_fill import fill_resume_template


def test_fill_resume_template_replaces_paragraph_and_table_placeholders(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "filled.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("{{summary}}")
    run.bold = True
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{skills}}"
    document.save(template_path)

    result = fill_resume_template(
        template_path=template_path,
        output_path=output_path,
        content_map={
            "summary": "具备 Java 后端和 RAG 项目经验",
            "skills": "Java / Spring Boot / Redis / RAG",
        },
    )

    filled = Document(output_path)
    assert filled.paragraphs[0].text == "具备 Java 后端和 RAG 项目经验"
    assert filled.paragraphs[0].runs[0].bold is True
    assert filled.tables[0].cell(0, 0).text == "Java / Spring Boot / Redis / RAG"
    assert result["status"] == "SUCCEEDED"
    assert result["placeholders"] == ["summary", "skills"]
    assert result["outputPath"] == str(output_path)


def test_fill_resume_template_reports_missing_placeholders(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "filled.docx"
    document = Document()
    document.add_paragraph("无占位符")
    document.save(template_path)

    result = fill_resume_template(
        template_path=template_path,
        output_path=output_path,
        content_map={"summary": "不会被替换"},
    )

    assert result["status"] == "SUCCEEDED"
    assert result["placeholders"] == []
    assert result["unusedKeys"] == ["summary"]
