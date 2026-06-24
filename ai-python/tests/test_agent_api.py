import os
import sys
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

AI_PYTHON_DIR = Path(__file__).resolve().parents[1]
if str(AI_PYTHON_DIR) not in sys.path:
    sys.path.insert(0, str(AI_PYTHON_DIR))

os.environ["RAG_STORE_BACKEND"] = "memory"

from app.main import app


class FakeResponse:
    def __init__(self, payload: dict | None = None, status_code: int = 200) -> None:
        self.payload = payload or {}
        self.status_code = status_code
        self.content = b"{}"

    def raise_for_status(self) -> None:
        if self.status_code >= 400:
            raise RuntimeError(f"HTTP {self.status_code}")

    def json(self) -> dict:
        return self.payload


class FakeJavaClient:
    calls: list[dict] = []

    def __init__(self, *args, **kwargs) -> None:
        self.headers = kwargs

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        return None

    def post(self, url: str, json: dict, headers: dict) -> FakeResponse:
        assert "/internal/rag/" not in url
        assert headers["X-Agent-Internal-Token"] == "agent-secret"
        FakeJavaClient.calls.append({"url": url, "json": json})
        if url.endswith("/api/internal/agent/tools/read"):
            if json["toolName"] == "agent_memory_retriever":
                return FakeResponse(
                    {
                        "taskId": json["taskId"],
                        "toolCallId": json["toolCallId"],
                        "toolName": json["toolName"],
                        "status": "SUCCEEDED",
                        "ownershipVerified": True,
                        "scope": "current_user_or_authorized",
                        "data": {
                            "memoryContext": [
                                {
                                    "memoryId": "agent-memory-1",
                                    "memoryType": "PREFERENCE",
                                    "namespace": "user_preference",
                                    "scope": "USER",
                                    "subjectKey": "answer_style",
                                    "summary": "用户偏好回答先给结论再列证据。",
                                    "score": 0.82,
                                }
                            ],
                            "memories": [],
                            "memoryCount": 1,
                        },
                        "diagnostics": {},
                        "retryable": False,
                    }
                )
            if json["toolName"] == "agent_memory_candidate_proposer":
                return FakeResponse(
                    {
                        "taskId": json["taskId"],
                        "toolCallId": json["toolCallId"],
                        "toolName": json["toolName"],
                        "status": "SUCCEEDED",
                        "ownershipVerified": True,
                        "scope": "current_user_or_authorized",
                        "data": {
                            "candidates": [
                                {
                                    "memoryType": "EPISODIC",
                                    "namespace": "agent_task",
                                    "scopeType": "USER",
                                    "subjectKey": "recent_task_insight",
                                    "content": "用户最近关注后端实习 JD 适配。",
                                    "summary": "后端实习 JD 适配关注 Java 和 RAG。",
                                    "confidence": 0.62,
                                    "importance": 0.56,
                                }
                            ],
                            "conflicts": [],
                            "provider": "fake",
                        },
                        "diagnostics": {},
                        "retryable": False,
                    }
                )
            if json["toolName"] == "web_search_probe":
                return FakeResponse(
                    {
                        "taskId": json["taskId"],
                        "toolCallId": json["toolCallId"],
                        "toolName": json["toolName"],
                        "status": "SUCCEEDED",
                        "ownershipVerified": True,
                        "scope": "current_user_or_authorized",
                        "data": {
                            "query": json["arguments"]["query"],
                            "retrievedAt": "2026-06-21T16:40:00+08:00",
                            "results": [
                                {
                                    "title": "公司技术趋势",
                                    "sourceUrl": "https://example.com/trend",
                                    "summary": "外部参考摘要",
                                    "score": 0.88,
                                    "confidence": "HIGH",
                                    "retrievedAt": "2026-06-21T16:40:00+08:00",
                                }
                            ],
                            "resultCount": 1,
                        },
                        "diagnostics": {},
                        "retryable": False,
                    }
                )
            return FakeResponse(
                {
                    "taskId": json["taskId"],
                    "toolCallId": json["toolCallId"],
                    "toolName": json["toolName"],
                    "status": "SUCCEEDED",
                    "ownershipVerified": True,
                    "scope": "current_user_or_authorized",
                    "data": {
                        "answer": "Redis 证据集中在缓存淘汰和持久化。",
                        "expandedQueries": ["Redis 缓存淘汰", "Redis 持久化"],
                        "evidences": [
                            {
                                "evidenceId": "material-12-1",
                                "title": "Redis 笔记",
                                "snippet": "正文片段不应进入 Observation 摘要",
                            }
                        ],
                        "diagnostics": {"candidateCount": 4},
                    },
                    "diagnostics": {},
                    "retryable": False,
                }
            )
        if url.endswith("/api/internal/agent/tools/mutation/execute"):
            return FakeResponse(
                {
                    "taskId": json["taskId"],
                    "toolCallId": json["toolCallId"],
                    "toolName": json["toolName"],
                    "status": "SUCCEEDED",
                    "ownershipVerified": True,
                    "scope": "current_user_or_authorized",
                    "data": {
                        "operationId": json["operationId"],
                        "status": "APPLIED_UNDOABLE",
                        "beforeSnapshotRef": "agent-operation-snapshot:snapshot-before-1",
                        "afterSnapshotRef": "agent-operation-snapshot:snapshot-after-1",
                        "undoDeadline": "2026-06-21T16:20:00+08:00",
                    },
                    "diagnostics": {},
                    "retryable": False,
                }
            )
        return FakeResponse({"accepted": True})


def test_agent_task_requires_internal_token(monkeypatch):
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)

    response = client.post(
        "/internal/agent/tasks",
        json={
            "taskId": "agent-task-1",
            "taskType": "pure_read_query",
            "input": {"goal": "Redis 学到了什么"},
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-1/events",
            "javaToolGatewayBaseUrl": "http://java",
            "threadId": "agent-task-1",
        },
    )

    assert response.status_code == 401
    assert response.json()["detail"] == "AGENT_INTERNAL_TOKEN_INVALID"


def test_agent_task_uses_java_gateway_and_callbacks(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-1",
            "taskType": "pure_read_query",
            "input": {
                "goal": "Redis 学到了什么",
                "topK": 3,
                "toolHints": ["rag_query_probe_non_persistent"],
                "metadataFilter": {"documentType": "markdown"},
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-1/events",
            "javaToolGatewayBaseUrl": "http://java",
            "threadId": "agent-task-1",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "COMPLETED"
    tool_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/api/internal/agent/tools/read")]
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert [call["json"]["toolName"] for call in tool_calls] == ["agent_memory_retriever", "rag_query_probe_non_persistent"]
    assert tool_calls[1]["json"]["arguments"]["metadataFilter"] == {"documentType": "markdown"}
    assert [event["json"]["eventType"] for event in event_calls] == [
        "TASK_STARTED",
        "TOOL_OBSERVATION",
        "TASK_COMPLETED",
    ]
    observation = event_calls[1]["json"]["toolCall"]["response"]
    assert observation["evidenceCount"] == 1
    assert "正文片段" not in str(observation)


def test_planning_task_requests_plan_review_after_memory_prefetch(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "input": {
                "goal": "分析后端实习 JD 适配度",
                "jobDescription": "要求 Java、Spring Boot、Redis 和 RAG 项目经验",
                "resumeText": "做过多模态 RAG 项目，熟悉 Java",
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
            "threadId": "agent-task-plan",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "WAITING_PLAN_REVIEW"
    tool_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/api/internal/agent/tools/read")]
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert [call["json"]["toolName"] for call in tool_calls] == ["agent_memory_retriever"]
    assert [event["json"]["eventType"] for event in event_calls] == ["TASK_STARTED", "REVIEW_REQUESTED"]
    assert event_calls[-1]["json"]["reviewRequest"]["reviewType"] == "PLAN"
    assert event_calls[-1]["json"]["draft"]["memoryContext"][0]["memoryId"] == "agent-memory-1"


def test_planning_resume_uses_java_gateway_and_requests_output_review(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks/agent-task-plan/resume",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "threadId": "agent-task-plan",
            "reviewType": "PLAN",
            "decision": "APPROVED",
            "decisionPayload": {"comment": "同意继续"},
            "input": {
                "goal": "分析后端实习 JD 适配度",
                "jobDescription": "要求 Java、Spring Boot、Redis 和 RAG 项目经验",
                "resumeText": "做过多模态 RAG 项目，熟悉 Java",
                "topK": 3,
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "WAITING_OUTPUT_REVIEW"
    tool_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/api/internal/agent/tools/read")]
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert [call["json"]["toolName"] for call in tool_calls] == [
        "agent_memory_retriever",
        "rag_query_probe_non_persistent",
        "agent_memory_candidate_proposer",
    ]
    assert [event["json"]["eventType"] for event in event_calls] == [
        "TOOL_OBSERVATION",
        "DRAFT_UPDATED",
        "REVIEW_REQUESTED",
    ]
    assert event_calls[-1]["json"]["reviewRequest"]["reviewType"] == "OUTPUT"
    assert event_calls[-1]["json"]["draft"]["alignment"][0]["status"] == "supported"
    assert event_calls[-1]["json"]["draft"]["pendingMemoryCandidates"][0]["summary"] == "后端实习 JD 适配关注 Java 和 RAG。"


def test_planning_resume_with_web_search_keeps_local_rag_flow(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks/agent-task-plan/resume",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "threadId": "agent-task-plan",
            "reviewType": "PLAN",
            "decision": "APPROVED",
            "decisionPayload": {"comment": "同意继续"},
            "input": {
                "goal": "分析后端实习 JD 适配度",
                "jobDescription": "要求 Java、Spring Boot、Redis 和 RAG 项目经验",
                "resumeText": "做过多模态 RAG 项目，熟悉 Java",
                "enableWebSearch": True,
                "toolHints": ["web_search_probe"],
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "WAITING_OUTPUT_REVIEW"
    tool_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/api/internal/agent/tools/read")]
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert [call["json"]["toolName"] for call in tool_calls] == [
        "agent_memory_retriever",
        "web_search_probe",
        "rag_query_probe_non_persistent",
        "agent_memory_candidate_proposer",
    ]
    assert [event["json"]["eventType"] for event in event_calls] == [
        "TOOL_OBSERVATION",
        "TOOL_OBSERVATION",
        "DRAFT_UPDATED",
        "REVIEW_REQUESTED",
    ]
    assert event_calls[2]["json"]["draft"]["webReferences"][0]["sourceUrl"] == "https://example.com/trend"


def test_planning_resume_can_fill_resume_template_draft(monkeypatch, tmp_path):
    import httpx

    template_path = tmp_path / "resume-template.docx"
    output_dir = tmp_path / "outputs"
    document = Document()
    document.add_paragraph("摘要：{{summary}}")
    document.add_paragraph("技能：{{skills}}")
    document.save(template_path)

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks/agent-task-plan/resume",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "threadId": "agent-task-plan",
            "reviewType": "PLAN",
            "decision": "APPROVED",
            "decisionPayload": {"comment": "同意继续"},
            "input": {
                "goal": "生成后端实习简历草稿",
                "jobDescription": "要求 Java、Redis 和 RAG 项目经验",
                "resumeText": "做过多模态 RAG 项目，熟悉 Java",
                "toolHints": ["resume_template_fill"],
                "resumeTemplatePath": str(template_path),
                "resumeTemplateOutputDir": str(output_dir),
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
        },
    )

    assert response.status_code == 200
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    draft = event_calls[1]["json"]["draft"]
    fill_result = draft["resumeTemplateFill"]
    assert fill_result["status"] == "SUCCEEDED"
    assert Path(fill_result["outputPath"]).exists()
    filled = Document(fill_result["outputPath"])
    assert "具备" in filled.paragraphs[0].text or "做过多模态" in filled.paragraphs[0].text
    assert "{{summary}}" not in filled.paragraphs[0].text


def test_planning_output_approval_with_save_intent_requests_crud_review(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks/agent-task-plan/resume",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "threadId": "agent-task-plan",
            "reviewType": "OUTPUT",
            "decision": "APPROVED",
            "decisionPayload": {"comment": "输出可保存"},
            "input": {
                "goal": "保存学习计划",
                "saveDraft": True,
                "toolHints": ["jd_learning_plan_save"],
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "WAITING_CRUD_REVIEW"
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert [event["json"]["eventType"] for event in event_calls] == ["MUTATION_PROPOSED"]
    review = event_calls[0]["json"]["reviewRequest"]
    assert review["reviewType"] == "CRUD"
    assert review["proposal"]["toolName"] == "jd_learning_plan_save"
    assert review["proposal"]["idempotencyKey"] == "jd_learning_plan_save-agent-task-plan-v1"


def test_planning_crud_approval_executes_java_mutation_gateway(monkeypatch):
    import httpx

    FakeJavaClient.calls = []
    monkeypatch.setenv("EVIDENCE_AGENT_INTERNAL_TOKEN", "agent-secret")
    client = TestClient(app)
    monkeypatch.setattr(httpx, "Client", FakeJavaClient)

    response = client.post(
        "/internal/agent/tasks/agent-task-plan/resume",
        headers={"X-Agent-Internal-Token": "agent-secret"},
        json={
            "taskId": "agent-task-plan",
            "taskType": "planning_task",
            "threadId": "agent-task-plan",
            "reviewType": "CRUD",
            "decision": "APPROVED",
            "decisionPayload": {
                "reviewId": "review-crud-agent-task-plan",
                "toolName": "jd_learning_plan_save",
                "idempotencyKey": "jd_learning_plan_save-agent-task-plan-v1",
                "comment": "同意保存",
            },
            "input": {
                "goal": "保存学习计划",
                "saveDraft": True,
                "toolHints": ["jd_learning_plan_save"],
            },
            "callbackUrl": "http://java/api/internal/agent/tasks/agent-task-plan/events",
            "javaToolGatewayBaseUrl": "http://java",
        },
    )

    assert response.status_code == 200
    assert response.json()["status"] == "COMPLETED"
    mutation_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/api/internal/agent/tools/mutation/execute")]
    event_calls = [call for call in FakeJavaClient.calls if call["url"].endswith("/events")]
    assert len(mutation_calls) == 1
    assert mutation_calls[0]["json"]["approvalId"] == "review-crud-agent-task-plan"
    assert mutation_calls[0]["json"]["toolName"] == "jd_learning_plan_save"
    assert [event["json"]["eventType"] for event in event_calls] == ["TOOL_OBSERVATION", "TASK_COMPLETED"]
    assert event_calls[0]["json"]["toolCall"]["toolType"] == "MUTATION"
    assert event_calls[-1]["json"]["final"]["operationStatus"] == "APPLIED_UNDOABLE"
