import base64
import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from app.schemas.resume_template import ResumeContentPatch, ResumePatchGenerationRequest, ResumeTemplatePreviewRequest
from rag.resume_template.docx_patch import (
    apply_resume_patches_to_docx,
    parse_resume_template_docx,
    validate_resume_patches,
)
from rag.resume_template.patch_generation import build_generation_prompt, generate_resume_patches, resume_patch_json_schema
from rag.resume_template.preview import build_resume_template_preview, render_pdf_preview


def docx_bytes(document: Document) -> bytes:
    """把测试 DOCX 文档保存为字节。"""
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_parse_resume_template_builds_bindings_and_replaces_single_run_table_cell():
    """解析字段绑定后，只替换表格单元格中的既有 w:t 文本节点。"""
    document = Document()
    document.add_heading("项目经历", level=1)
    document.add_paragraph("多模态 RAG 学习证据平台，负责 FastAPI 检索接口。")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "技能：Python、Spring Boot、React"
    content = docx_bytes(document)

    parsed = parse_resume_template_docx(content, "resume.docx", template_id="tpl-1")
    table_field = next(field for field in parsed.fields if field.locationRefs[0].containerType == "table_cell")
    patch = ResumeContentPatch(
        fieldId=table_field.fieldId,
        sourceTextHash=table_field.sourceTextHash,
        newText="技能：Python、FastAPI、Spring Boot、React",
        rewriteReason="根据 JD 强化 FastAPI 表述",
        evidenceIds=["e1"],
        confidence=0.9,
        riskFlags=["NONE"],
        status="CONFIRMED",
    )

    exported = apply_resume_patches_to_docx(
        content,
        "resume.docx",
        "tpl-1",
        1,
        parsed.fields,
        [patch],
        allowed_evidence_ids=["e1"],
    )
    exported_bytes = base64.b64decode(exported.fileBase64)
    result_doc = Document(io.BytesIO(exported_bytes))

    assert exported.appliedPatchCount == 1
    assert exported.layoutValidation["status"] == "PASSED"
    assert result_doc.tables[0].cell(0, 0).text == "技能：Python、FastAPI、Spring Boot、React"


def test_validate_resume_patches_rejects_unknown_hash_and_markup_injection():
    """补丁校验拒绝未知字段、hash 不匹配和 HTML/XML/Markdown 注入。"""
    document = Document()
    document.add_paragraph("项目：RAG 检索平台")
    parsed = parse_resume_template_docx(docx_bytes(document), "resume.docx", template_id="tpl-2")
    field = parsed.fields[0]
    patches = [
        ResumeContentPatch(
            fieldId="missing",
            sourceTextHash=field.sourceTextHash,
            newText="无效字段",
            rewriteReason="测试未知字段",
            evidenceIds=[],
            confidence=0.1,
            riskFlags=["LOW_CONFIDENCE"],
            status="DRAFT",
        ),
        ResumeContentPatch(
            fieldId=field.fieldId,
            sourceTextHash="bad-hash-value-0000",
            newText="<w:t>注入</w:t>",
            rewriteReason="测试注入",
            evidenceIds=[],
            confidence=0.1,
            riskFlags=["INJECTION_RISK"],
            status="DRAFT",
        ),
    ]

    result = validate_resume_patches("tpl-2", 1, parsed.fields, patches, allowed_evidence_ids=[])

    assert not result.patches
    assert "不存在对应字段" in "；".join(result.validationErrors)
    assert "原文字段 hash 不匹配" in "；".join(result.validationErrors)
    assert "注入风险" in "；".join(result.validationErrors)


def test_apply_rejects_cross_run_field_to_protect_layout():
    """跨 run 字段首版拒绝自动替换，避免清空 follower run 改变结构语义。"""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("项目：")
    paragraph.add_run("RAG 平台")
    content = docx_bytes(document)
    parsed = parse_resume_template_docx(content, "resume.docx", template_id="tpl-3")

    assert parsed.fields == []


def test_export_rejects_when_source_hash_changed():
    """导出时如果当前 DOCX 字段文本已变化，返回版本冲突式错误。"""
    document = Document()
    document.add_paragraph("原始项目经历")
    content = docx_bytes(document)
    parsed = parse_resume_template_docx(content, "resume.docx", template_id="tpl-4")
    field = parsed.fields[0]
    patch = ResumeContentPatch(
        fieldId=field.fieldId,
        sourceTextHash=field.sourceTextHash,
        newText="确认后的项目经历",
        rewriteReason="用户确认改写",
        evidenceIds=["e1"],
        confidence=0.9,
        riskFlags=["NONE"],
        status="CONFIRMED",
    )
    changed = Document(io.BytesIO(content))
    changed.paragraphs[0].runs[0].text = "用户已经手动改过"

    with pytest.raises(ValueError, match="hash 已变化"):
        apply_resume_patches_to_docx(docx_bytes(changed), "resume.docx", "tpl-4", 1, parsed.fields, [patch], allowed_evidence_ids=["e1"])


def test_unsupported_hyperlink_region_blocks_export(tmp_path: Path):
    """包含超链接关系的 DOCX 会进入 unsupportedRegions，并拒绝导出。"""
    document = Document()
    document.add_paragraph("项目经历")
    content = add_hyperlink_xml(docx_bytes(document), tmp_path)
    parsed = parse_resume_template_docx(content, "resume.docx", template_id="tpl-5")

    assert any("超链接" in item for item in parsed.unsupportedRegions)
    with pytest.raises(ValueError, match="复杂结构"):
        apply_resume_patches_to_docx(content, "resume.docx", "tpl-5", 1, parsed.fields, [])


def test_export_allows_static_picture_when_patch_targets_plain_field(tmp_path: Path):
    """模板包含头像或图片时，只要补丁字段不在图片区域，仍允许导出。"""
    document = Document()
    document.add_paragraph("项目：多模态 RAG 学习证据平台")
    image_path = tmp_path / "avatar.png"
    image_path.write_bytes(base64.b64decode(TINY_PNG_BASE64))
    document.add_picture(str(image_path), width=Inches(0.25))
    content = docx_bytes(document)
    parsed = parse_resume_template_docx(content, "resume.docx", template_id="tpl-image")
    field = parsed.fields[0]
    patch = ResumeContentPatch(
        fieldId=field.fieldId,
        sourceTextHash=field.sourceTextHash,
        newText="项目：多模态 RAG 学习证据平台，负责模板导出链路",
        rewriteReason="根据 JD 强化职责表述",
        evidenceIds=["e1"],
        confidence=0.9,
        riskFlags=["NONE"],
        status="CONFIRMED",
    )

    assert any("绘图或图片" in item for item in parsed.unsupportedRegions)
    exported = apply_resume_patches_to_docx(
        content,
        "resume.docx",
        "tpl-image",
        1,
        parsed.fields,
        [patch],
        allowed_evidence_ids=["e1"],
    )
    exported_bytes = base64.b64decode(exported.fileBase64)
    result_doc = Document(io.BytesIO(exported_bytes))

    assert exported.layoutValidation["status"] == "PASSED"
    assert result_doc.paragraphs[0].text == "项目：多模态 RAG 学习证据平台，负责模板导出链路"


def test_local_patch_generation_returns_strict_schema_and_safe_drafts():
    """无模型环境下返回严格 schema 和保留原文的安全草稿。"""
    document = Document()
    document.add_paragraph("技能：RAG、BM25、FastAPI")
    parsed = parse_resume_template_docx(docx_bytes(document), "resume.docx", template_id="tpl-6")
    response = generate_resume_patches(
        ResumePatchGenerationRequest.model_validate({
            "templateId": "tpl-6",
            "version": 1,
            "jobDescription": "需要 RAG 和 FastAPI 经验",
            "fields": [field.model_dump() for field in parsed.fields],
            "evidenceCandidates": [],
            "provider": "local",
        })
    )
    schema = resume_patch_json_schema()

    assert schema["additionalProperties"] is False
    assert schema["properties"]["patches"]["items"]["additionalProperties"] is False
    assert response.provider == "local"
    assert response.patches[0].newText == parsed.fields[0].sourceText


def test_patch_generation_prompt_includes_uploaded_resume_summary():
    """字段补丁提示词必须包含用户已上传简历摘要，避免只按 JD 改模板。"""
    document = Document()
    document.add_paragraph("项目：多模态 RAG 平台")
    parsed = parse_resume_template_docx(docx_bytes(document), "resume.docx", template_id="tpl-summary")
    request = ResumePatchGenerationRequest.model_validate({
        "templateId": "tpl-summary",
        "version": 1,
        "jobDescription": "需要 Spring Boot 和 RAG 项目经验",
        "resumeText": "已上传简历摘要：负责 FastAPI 检索、React 模板确认和 Java 联调。",
        "fields": [field.model_dump() for field in parsed.fields],
        "evidenceCandidates": [],
        "provider": "local",
    })

    prompt = build_generation_prompt(request)

    assert "用户已上传简历摘要" in prompt
    assert "FastAPI 检索" in prompt


def test_preview_maps_pdf_text_block_to_relative_rect():
    """图片预览把 PDF 文本块匹配回字段，并保证坐标在 0..1 范围。"""
    document = Document()
    document.add_paragraph("项目：多模态 RAG 学习证据平台")
    document.add_paragraph("未映射字段内容")
    parsed = parse_resume_template_docx(docx_bytes(document), "resume.docx", template_id="tpl-preview")
    request = {
        "templateId": "tpl-preview",
        "version": 1,
        "filename": "resume.docx",
        "fileBase64": base64.b64encode(docx_bytes(document)).decode("ascii"),
        "fields": [field.model_dump() for field in parsed.fields],
    }

    response = render_pdf_preview(
        FakeFitzModule(),
        Path("fake.pdf"),
        ResumeTemplatePreviewRequest.model_validate(request),
        [],
    )

    assert response.previewStatus == "PARTIAL"
    assert response.pages[0].imageMimeType == "image/png"
    assert len(response.regions) == 1
    rect = response.regions[0].rect
    assert 0 <= rect.x <= 1
    assert 0 <= rect.y <= 1
    assert 0 < rect.width <= 1
    assert 0 < rect.height <= 1
    assert response.unmappedFields


def test_preview_returns_field_sketch_when_libreoffice_missing(monkeypatch):
    """LibreOffice 不可用时返回字段草图预览，图片确认流程仍可使用。"""
    pytest.importorskip("fitz")
    document = Document()
    document.add_paragraph("项目：RAG 平台")
    parsed = parse_resume_template_docx(docx_bytes(document), "resume.docx", template_id="tpl-no-preview")
    monkeypatch.setattr("rag.resume_template.preview.convert_with_libreoffice", lambda *args, **kwargs: None)

    response = build_resume_template_preview(ResumeTemplatePreviewRequest.model_validate({
        "templateId": "tpl-no-preview",
        "version": 1,
        "filename": "resume.docx",
        "fileBase64": base64.b64encode(docx_bytes(document)).decode("ascii"),
        "fields": [field.model_dump() for field in parsed.fields],
    }))

    assert response.previewStatus == "PARTIAL"
    assert response.pages
    assert response.regions
    assert response.regions[0].previewStatus == "PARTIAL"
    assert response.regions[0].confidence == 0.36
    assert response.pages[0].imageMimeType == "image/png"
    assert response.pages[0].imageBase64
    assert response.unmappedFields == []
    assert "LibreOffice" in "；".join(response.warnings)


class FakePixmap:
    """提供测试用 PNG 字节。"""

    def tobytes(self, image_type: str) -> bytes:
        return b"fake-png"


class FakePage:
    """模拟 PyMuPDF 页面对象。"""

    rect = type("Rect", (), {"width": 200, "height": 100})()

    def get_pixmap(self, matrix, alpha: bool):
        return FakePixmap()

    def get_text(self, mode: str):
        return [(20, 10, 180, 30, "项目：多模态 RAG 学习证据平台", 0, 0)]


class FakeDocument:
    """模拟 PyMuPDF 文档上下文。"""

    def __enter__(self):
        return [FakePage()]

    def __exit__(self, exc_type, exc, traceback):
        return False


class FakeFitzModule:
    """模拟 PyMuPDF 模块接口。"""

    def Matrix(self, x: int, y: int):
        return (x, y)

    def open(self, path: str):
        return FakeDocument()


def add_hyperlink_xml(content: bytes, tmp_path: Path) -> bytes:
    """直接在 document.xml 注入 hyperlink，用于测试复杂结构检测。"""
    source = tmp_path / "source.docx"
    target = tmp_path / "target.docx"
    source.write_bytes(content)
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                text = text.replace("<w:t>项目经历</w:t>", "<w:hyperlink><w:r><w:t>项目经历</w:t></w:r></w:hyperlink>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    return target.read_bytes()


TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)
